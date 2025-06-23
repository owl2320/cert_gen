import customtkinter as ctk
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox,PhotoImage,filedialog
from tkinterdnd2 import TkinterDnD, DND_FILES
from bijoy2unicode.converter import Unicode
from docx import Document
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
import os
import re
import pandas as pd

def gen_certificate(name, data, template_path, output_folder):
    try:
        doc = DocxTemplate(template_path)
        doc.render(data)
        doc.save(os.path.join(output_folder, name))
    except Exception as e:
        messagebox.showerror("Error", f"Error generating certificate '{name}': {e}")

def merge_documents(output_folder, merged_name="All_Certificates_Merged.docx"):
    try:
        files = sorted(
            [f for f in os.listdir(output_folder) if f.endswith(".docx") and f != merged_name],
            key=lambda x: int(re.search(r'\d+', x).group()) if re.search(r'\d+', x) else x
        )
        if not files:
            raise FileNotFoundError("No .docx files found to merge.")

        master = Document(os.path.join(output_folder, files[0]))
        composer = Composer(master)

        for file in files[1:]:
            composer.append(Document(os.path.join(output_folder, file)))
        composer.save(os.path.join(output_folder, merged_name))
    except Exception as e:
        messagebox.showerror("Error", f"Failed to merge documents: {e}")

def gen_certificate_csv(csv_path, template_path, output_folder, progress_callback):
    """generate_certificates csv"""
    try:
        csv_data = pd.read_csv(csv_path)
    except Exception as e:
        messagebox.showerror("Error", f"Error reading CSV file: {e}")
        return

    total = len(csv_data)
    for idx, info in csv_data.iterrows():
        try:
            data = {
                'id': info["Teacher's Training ID"],
                'name': str(info["Teacher's Name"]).title(),
                'position': str(info["Teacher's Designation"]).title(),
                'school_name': str(info["Name of Institution"]).title()
            }
            gen_certificate(f'{idx + 1}.docx', data, template_path, output_folder)
            progress_callback(idx + 1, total)
        except Exception as e:
            messagebox.showerror("Error", f"Error on row {idx + 1}: {e}")

def gen_certificate_docx(docx_path, template_path, output_folder, progress_callback):
    """generate certificates docx"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        messagebox.showerror("Error", f"Error reading DOCX file: {e}")
        return

    table = doc.tables[0]
    total = len(table.rows)
    for idx, row in enumerate(table.rows):
        try:
            cells = [cell.text.strip() for cell in row.cells]
            name_lines = cells[1].split('\n')
            school_name = ''.join(name_lines[2:]).strip()
            data = {
                'id': cells[0],
                'name': name_lines[0],
                'position': name_lines[1] if len(name_lines) > 1 else '',
                'school_name': re.sub(r'[0-9]', '', school_name).rstrip(", ")
            }
            gen_certificate(f'{idx + 1}.docx', data, template_path, output_folder)
            progress_callback(idx + 1, total)
        except Exception as e:
            messagebox.showerror("Error", f"Error on row {idx + 1}: {e}")

def gen_certificate_docx_bijoy(docx_path, template_path, output_folder, progress_callback):
    """generate certificates from bijoy"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        messagebox.showerror("Error", f"Error reading DOCX file: {e}")
        return

    unicode = Unicode()
    col1, col2, col3 = [], [], []

    for table in doc.tables:
        for row in table.rows:
            try:
                col1.append(unicode.convertBijoyToUnicode(row.cells[0].text.strip()))
                col2.append(unicode.convertBijoyToUnicode(row.cells[1].text.strip()))
                col3.append(unicode.convertBijoyToUnicode(row.cells[2].text.strip()))
            except IndexError:
                continue

    col1, col2, col3 = col1[2:], col2[2:], col3[2:]
    total = len(col2)

    for i in range(total):
        try:
            name_pos = col2[i].split("\n")
            school_lines = col3[i].split("\n")
            data = {
                "id": col1[i],
                "name": name_pos[0],
                "position": name_pos[1] if len(name_pos) > 1 else name_pos[0],
                "school_name": school_lines[0] if school_lines else ""
            }
            gen_certificate(f'{i + 1}.docx', data, template_path, output_folder)
            progress_callback(i + 1, total)
        except Exception as e:
            messagebox.showerror("Error", f"Error on row {i + 1}: {e}")

def drop(event, obj):
    path = event.data
    obj.delete(0, 'end')
    obj.insert(0, path.lstrip('{').rstrip('}'))

def browse_input(e):
    input_path = filedialog.askopenfile(title="Select File",
                                        filetypes=[("Docx", ".docx"), ("CSV", ".csv"), ("Excel", ".xlsx")])
    if input_path:
        e.delete(0, 'end')
        e.insert(0, input_path.name)

def browse_template(e):
    input_path = filedialog.askopenfile(title="Select File",
                                        filetypes=[("Docx", ".docx")])
    if input_path:
        e.delete(0, 'end')
        e.insert(0, input_path.name)

def browse_folder(e):
    input_path = filedialog.askdirectory()
    if input_path:
        e.delete(0, 'end')
        e.insert(0, input_path)

def change_appearence(mode, win, input_label, temp_label,merge_files,bijoy_uni,output_folder,progress_label):
    if mode:
        win.config(bg='#F0F0F0')
        input_label.configure(text_color='black')
        temp_label.configure(text_color='black')
        merge_files.configure(text_color="black",border_color='black')
        bijoy_uni.configure(text_color="black",border_color='black')
        output_folder.configure(text_color='black')
        progress_label.configure(text_color='black')
    else:
        win.config(bg='#242424')
        input_label.configure(text_color='white')
        temp_label.configure(text_color='white')
        merge_files.configure(text_color="white",border_color='#F0F0F0')
        bijoy_uni.configure(text_color="white",border_color='#F0F0F0')
        output_folder.configure(text_color='white')
        progress_label.configure(text_color='white')

def show_about():
    messagebox.showinfo("About The Developer", "This was made by Abrar Jawad Al Tasin."
                                               " It is designed to help my father in his work."
                                               " If you're not him and having problems, ask him for help.")

ctk.set_appearance_mode('Light')

class Main:
    def __init__(self):
        self.win = TkinterDnD.Tk()
        self.win.title("Certificate Generator")
        self.win.geometry("500x300")
        self.win.resizable(False,False)

        #check for wheather on or not
        self.merge_var = tk.BooleanVar()
        self.bijoy_var = tk.BooleanVar()

        #file paths
        self.file_path = tk.StringVar()
        self.template_path = tk.StringVar()
        self.output_path = tk.StringVar()

    def update_progress(self,current, total):
        percent = int((current / total) * 100)
        self.progress["value"] = percent
        self.progress_label.configure(text=f"{current}/{total} certificates generated")
        self.win.update_idletasks()

    def generate(self):
        """generate all the certificates"""
        try:
            self.progress.configure(value=0)
            self.progress_label.configure(text="Processing...")

            if self.file_path.get().endswith('.docx'):
                if self.bijoy_var.get():
                    gen_certificate_docx_bijoy(self.file_path.get()
                                               ,self.template_path.get(),
                                               self.output_path.get(),
                                               self.update_progress)
                else:
                    gen_certificate_docx(self.file_path.get(),
                                         self.template_path.get(),
                                         self.output_path.get(),
                                         self.update_progress)
            if self.file_path.get().endswith('.csv'):
                gen_certificate_csv(self.file_path.get(),
                                         self.template_path.get(),
                                         self.output_path.get(),
                                         self.update_progress)

            if self.merge_var.get():
                merge_documents(self.output_path.get())

            if not (self.file_path.get() or self.template_path.get() or self.output_path.get()):
                messagebox.showerror("Error", "Not selected all the file paths.")
            else:
                messagebox.showinfo('Success',"All Certificates Generated Successfully.")
            self.progress.configure(value=0)
            self.progress_label.configure(text="")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def menubar(self):
        menubar = tk.Menu(self.win)

        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Mode", menu=file_menu)
        file_menu.add_command(label="Light", command=lambda: change_appearence(1, self.win,
                                                                               self.input_label,
                                                                               self.temp_label,
                                                                               self.merge_files,
                                                                               self.bijoy_to_uni,
                                                                               self.output_label,
                                                                               self.progress_label))
        file_menu.add_command(label="Dark", command=lambda: change_appearence(0, self.win,
                                                                              self.input_label,
                                                                              self.temp_label,
                                                                              self.merge_files,
                                                                              self.bijoy_to_uni,
                                                                              self.output_label,
                                                                              self.progress_label))

        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="About", command=show_about)

        self.win.config(menu=menubar)

    def input_file(self):
        self.input_label = ctk.CTkLabel(self.win, text="Input File:", font=("Calibri", 18, "bold"))
        self.input_label.place(x=5, y=5)

        self.input_entry = ctk.CTkEntry(self.win,
                                        width=300,
                                        corner_radius=1,
                                        border_width=1,
                                        textvariable=self.file_path)
        self.input_entry.place(x=120, y=5)
        self.input_entry.drop_target_register(DND_FILES)
        self.input_entry.dnd_bind('<<Drop>>', lambda event: drop(event, self.input_entry))

        self.input_browse = ctk.CTkButton(self.win, text="Browse..", width=10, corner_radius=3,
                                          command=lambda: browse_input(self.input_entry),
                                          fg_color='white', text_color='black', hover_color='#00CCFF')
        self.input_browse.place(x=430, y=5)

    def template_file(self):
        self.temp_label = ctk.CTkLabel(self.win, text="Template File:", font=("Calibri", 18, "bold"))
        self.temp_label.place(x=5, y=40)

        self.temp_entry = ctk.CTkEntry(self.win,
                                       width=300,
                                       corner_radius=1,
                                       border_width=1,
                                       textvariable=self.template_path)
        self.temp_entry.place(x=120, y=40)
        self.temp_entry.drop_target_register(DND_FILES)
        self.temp_entry.dnd_bind('<<Drop>>', lambda event: drop(event, self.temp_entry))

        self.temp_browse = ctk.CTkButton(self.win, text="Browse..", width=10, corner_radius=3,
                                         command=lambda: browse_template(self.temp_entry),
                                         fg_color='white', text_color='black', hover_color='#00CCFF')
        self.temp_browse.place(x=430, y=40)

    def output_folder(self):
        self.output_label = ctk.CTkLabel(self.win, text="Output Folder:", font=("Calibri", 18, "bold"))
        self.output_label.place(x=5, y=90)

        self.output_entry = ctk.CTkEntry(self.win,
                                         width=300,
                                         corner_radius=1,
                                         border_width=1,
                                         textvariable=self.output_path)
        self.output_entry.place(x=120, y=90)
        self.output_entry.drop_target_register(DND_FILES)
        self.output_entry.dnd_bind('<<Drop>>', lambda event: drop(event, self.output_entry))

        self.output_browse = ctk.CTkButton(self.win, text="Browse..", width=10, corner_radius=3,
                                           command=lambda: browse_folder(self.output_entry),
                                           fg_color='white', text_color='black', hover_color='#00CCFF')
        self.output_browse.place(x=430, y=90)

    def change_progress_bar(self,value):
        self.progress.configure(value=value)

    def add_checks(self):
        self.merge_files = ctk.CTkCheckBox(self.win, text="Merge", variable=self.merge_var)
        self.merge_files.place(x=10, y=140)

        self.bijoy_to_uni = ctk.CTkCheckBox(self.win, text="Bijoy to Unicode", variable=self.bijoy_var)
        self.bijoy_to_uni.place(x=100, y=140)

        ctk.CTkButton(self.win, text="Generate",
                      corner_radius=3,
                      width=50, height=20,
                      command=lambda:self.generate(),
                      font=("Calibri", 20)).place(x=290, y=140)

    def main_loop(self):
        self.input_file()
        self.template_file()
        self.menubar()
        self.add_checks()
        self.output_folder()
        #progress bar
        self.progress = ttk.Progressbar(self.win, length=480, mode='determinate')
        self.progress.place(x=10,y=220)
        #label
        self.progress_label = ctk.CTkLabel(self.win,text="")
        self.progress_label.place(x=10,y=180)

        self.win.mainloop()

if __name__ == '__main__':
    win = Main()
    win.main_loop()